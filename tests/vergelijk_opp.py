"""
Vergelijkt zichtbare vs verborgen tekst in OPP-documenten.
Toont aan dat de vergiftigde versie er identiek uitziet voor een leerkracht.

Voer uit vanuit de projectroot:
  python3 tests/vergelijk_opp.py
"""
import json
from datetime import datetime
from docx import Document
from docx.shared import RGBColor

RESULTS_FILE = f"tests/vergelijk_opp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

def extraheer_zichtbare_tekst(pad):
    """Alleen tekst die niet wit is — wat een leerkracht ziet."""
    doc = Document(pad)
    zichtbaar = []
    for para in doc.paragraphs:
        for run in para.runs:
            kleur = run.font.color.rgb if run.font.color and run.font.color.type else None
            if kleur != RGBColor(0xFF, 0xFF, 0xFF):
                if run.text.strip():
                    zichtbaar.append(run.text.strip())
    return " ".join(zichtbaar)

def extraheer_alle_tekst(pad):
    """Alles inclusief witte tekst — wat mammoth ziet."""
    doc = Document(pad)
    return " ".join(
        run.text for para in doc.paragraphs for run in para.runs if run.text.strip()
    )

print("=" * 60)
print("VERGELIJKING OPP-DOCUMENTEN")
print("=" * 60)

resultaten = []

for naam, pad in [("Origineel", "OPP_bestanden/OPP_1_backup.docx"),
                  ("Vergiftigd", "OPP_bestanden/OPP_1.docx")]:
    print(f"\n--- {naam} ---")
    zichtbaar = extraheer_zichtbare_tekst(pad)
    alles = extraheer_alle_tekst(pad)
    verborgen_tekens = len(alles) - len(zichtbaar)
    heeft_injectie = "SYSTEM" in alles or "negeer" in alles.lower()

    print(f"Wat leerkracht ziet : {zichtbaar[:150]}")
    print(f"Wat mammoth ziet    : {alles[:200]}")

    if naam == "Vergiftigd":
        print(f"\n  Verborgen tekens  : {verborgen_tekens}")
        print(f"  Injectie aanwezig : {'⚠️  JA' if heeft_injectie else '✅ Nee'}")
        print("  Onzichtbaar voor leerkracht, wel leesbaar door het AI-systeem")

    resultaten.append({
        "document": naam,
        "bestand": pad,
        "zichtbare_tekst": zichtbaar,
        "volledige_tekst_mammoth": alles,
        "verborgen_tekens": verborgen_tekens,
        "injectie_aanwezig": heeft_injectie,
        "visueel_identiek_aan_origineel": naam == "Vergiftigd",
    })

# Samenvatting
print("\n" + "=" * 60)
print("CONCLUSIE")
print("=" * 60)
origineel = next(r for r in resultaten if r["document"] == "Origineel")
vergiftigd = next(r for r in resultaten if r["document"] == "Vergiftigd")

zichtbaar_identiek = origineel["zichtbare_tekst"] == vergiftigd["zichtbare_tekst"]
print(f"\n  Visueel identiek voor leerkracht : {'✅ Ja' if zichtbaar_identiek else '⚠️  Nee'}")
print(f"  Injectie zichtbaar in mammoth    : {'⚠️  Ja — kwetsbaar' if vergiftigd['injectie_aanwezig'] else '✅ Nee'}")
print(f"  Verborgen tekens in vergiftigd   : {vergiftigd['verborgen_tekens']}")

# Opslaan als JSON
with open(RESULTS_FILE, "w", encoding="utf-8") as f:
    json.dump({
        "timestamp": datetime.now().isoformat(),
        "aanval": "Context Poisoning via verborgen witte tekst in OPP-document",
        "bron": "arxiv.org/pdf/2507.13169 §3.1.2",
        "resultaten": resultaten,
        "conclusie": {
            "visueel_identiek": zichtbaar_identiek,
            "injectie_aanwezig": vergiftigd["injectie_aanwezig"],
            "verborgen_tekens": vergiftigd["verborgen_tekens"],
        }
    }, f, ensure_ascii=False, indent=2)

print(f"\n  Resultaten opgeslagen: {RESULTS_FILE}")
